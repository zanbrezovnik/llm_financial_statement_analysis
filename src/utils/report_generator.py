"""
Creates Word documents (.docx) for chat transcripts and extraction logs.

This module uses the `python-docx` library to create formatted .docx files.
It includes functionality to parse simple Markdown (bold, italics, lists, tables)
from the chatbot's string output into formatted content in the Word document.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
import logging
from datetime import datetime
from typing import List, Dict, Any, Tuple
import re

logger = logging.getLogger(__name__)

def generate_extraction_log(
    log_entries: List[Dict[str, Any]], 
    output_filepath: str
) -> None:
    """Creates a .docx log file for the table extraction process."""
    doc = Document()
    doc.add_heading('PDF Table Extraction Log', level=1)
    doc.add_paragraph(f"Report generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("-" * 30)

    if not log_entries:
        doc.add_paragraph("No PDF files were processed.")
    else:
        for entry in log_entries:
            doc.add_heading(f"File: {entry.get('pdf_filename', 'Unknown File')}", level=2)
            p = doc.add_paragraph()
            p.add_run("Status: ").bold = True
            p.add_run(str(entry.get('status', 'N/A')))
            
            if 'processing_time_seconds' in entry:
                p_time = doc.add_paragraph()
                p_time.add_run("Processing Time: ").bold = True
                p_time.add_run(f"{entry['processing_time_seconds']:.2f} seconds")

            p_msg = doc.add_paragraph()
            p_msg.add_run("Details: ").bold = True
            p_msg.add_run(str(entry.get('message', 'No details.')))

            p_tables = doc.add_paragraph()
            p_tables.add_run("Attempted/Extracted Tables: ").bold = True
            tables_list = entry.get('extracted_tables', [])
            if tables_list:
                 p_tables.add_run(", ".join(tables_list) if isinstance(tables_list, list) else str(tables_list))
            else:
                p_tables.add_run("None.")
            doc.add_paragraph("-" * 20)
            
    try:
        doc.save(output_filepath)
        logger.info(f"Extraction log saved to: {output_filepath}")
    except Exception as e:
        logger.error(f"Failed to save extraction log to {output_filepath}: {e}")
        print(f"Error: Failed to save extraction log to {output_filepath}: {e}")

def _is_markdown_table_line(line: str) -> bool:
    """Checks if a line is part of a Markdown table (e.g., | a | b |)."""
    return line.strip().startswith("|") and line.strip().endswith("|")

def _is_markdown_table_separator(line: str) -> bool:
    """Checks if a line is a Markdown table separator (e.g., |---|---:|)."""
    stripped_line = line.strip()
    if not (_is_markdown_table_line(stripped_line) and stripped_line.count('|') >= 2):
        return False
    
    parts = [part.strip() for part in stripped_line.strip('|').split('|')]
    if not parts: # Should not happen if count('|') >= 2 but good for robustness
        return False

    # Check if content between pipes is only hyphens, optionally surrounded by colons
    for part in parts:
        if not part: # Can happen with || in separator, treat as valid
            continue
        # Remove optional colons at the start and end for the hyphen check
        core_part = part
        if core_part.startswith(':'):
            core_part = core_part[1:]
        if core_part.endswith(':'):
            core_part = core_part[:-1]
        
        if not core_part or not all(c == '-' for c in core_part):
            return False
    return True

def _parse_markdown_table_row(line: str) -> List[str]:
    """Splits a Markdown table row into a list of cells."""
    return [cell.strip() for cell in line.strip().strip('|').split('|')]

def _add_formatted_text(paragraph, text):
    """Adds text to a paragraph, handling **bold** and *italic* formatting."""
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

def generate_chat_transcript(
    chat_history: List[Tuple[str, str]], 
    loaded_pdfs: List[str], 
    output_filepath: str
) -> None:
    """
    Creates a .docx transcript of the chat.
    It formats the conversation and handles simple Markdown like tables and lists.
    """
    doc = Document()
    doc.add_heading('Chatbot Session Transcript', level=1)
    
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    context_pdfs = ', '.join(loaded_pdfs) if loaded_pdfs else "None"
    doc.add_paragraph(f"Context PDF(s) Loaded: {context_pdfs}")
    doc.add_paragraph("-" * 30)

    for user_query, chatbot_response in chat_history:
        # Format the user's query.
        p_user = doc.add_paragraph()
        p_user.add_run("User: ").bold = True
        run_user_query = p_user.add_run(user_query)
        run_user_query.font.color.rgb = RGBColor(0xC0, 0x00, 0x00) # Darker Red

        # Format the chatbot's response.
        p_chatbot_intro = doc.add_paragraph()
        p_chatbot_intro.add_run("Chatbot:").bold = True
        
        # This block parses the chatbot's Markdown response line-by-line.
        response_lines = chatbot_response.splitlines()
        i = 0
        while i < len(response_lines):
            line = response_lines[i]
            
            # Check for and process a Markdown table.
            if _is_markdown_table_line(line) and \
               i + 1 < len(response_lines) and _is_markdown_table_separator(response_lines[i+1]):
                
                table_rows = []
                # Header
                table_rows.append(_parse_markdown_table_row(line))
                i += 2 # Move past header and separator
                
                # Data rows
                while i < len(response_lines) and _is_markdown_table_line(response_lines[i]):
                    table_rows.append(_parse_markdown_table_row(response_lines[i]))
                    i += 1

                # Create Word table
                if table_rows:
                    num_cols = len(table_rows[0])
                    word_table = doc.add_table(rows=len(table_rows), cols=num_cols, style='Table Grid')
                    for r, row_data in enumerate(table_rows):
                        for c, cell_text in enumerate(row_data):
                            if c < num_cols:
                                word_table.cell(r, c).text = cell_text
                    # Set header bold
                    for cell in word_table.rows[0].cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                continue # Continue to next line after table processing

            # Check for and process a list item.
            list_match = re.match(r"^\s*([*\-+]|\d+\.)\s+", line)
            if list_match:
                p = doc.add_paragraph(style='List Bullet' if not list_match.group(1).endswith('.') else 'List Number')
                _add_formatted_text(p, line[list_match.end():])
            # Otherwise, it's just a regular paragraph.
            else:
                p = doc.add_paragraph()
                _add_formatted_text(p, line)
            
            i += 1
        
        doc.add_paragraph() # Add space between messages.

    try:
        doc.save(output_filepath)
        logger.info(f"Chat transcript successfully saved to {output_filepath}")
    except Exception as e:
        logger.error(f"Failed to save chat transcript to {output_filepath}: {e}", exc_info=True)
        print(f"Error: Failed to save chat transcript to {output_filepath}: {e}") 